import os
import json
from docx import Document
import requests
from io import BytesIO
from minio import Minio
import io


# Инициализация MinIO клиента
def get_minio_client():
    from scrapy.utils.project import get_project_settings
    settings = get_project_settings()
    minio_config = settings.get('MINIO_CONFIG', {})

    return Minio(
        minio_config['endpoint'],
        access_key=minio_config['access_key'],
        secret_key=minio_config['secret_key'],
        secure=minio_config['secure']
    ), minio_config['bucket_name']


# Основной метод для извлечения текста из документа
def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        text_data = []

        # Извлекаем текст из параграфов
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                text_data.append(para_text)

        # Извлекаем текст из таблиц
        table_data = extract_table_data(file_path)
        for table_row in table_data:
            text_data.append(" | ".join(table_row))

        # Объединяем все извлеченные данные в одну строку
        merge_text = " ".join(text_data)
        return merge_text

    except Exception as e:
        print(f"Ошибка при обработке документа: {e}")
        return ""


# Метод для извлечения данных из таблиц
def extract_table_data(docx_path):
    doc = Document(docx_path)
    table_data = []

    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if '\n' in cell_text:
                    cell_text = " ".join(cell_text.split('\n')).strip()
                row_data.append(cell_text)
            table_data.append(row_data)

    return table_data


# Функция для скачивания документа по URL
def download_docx(url):
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        docx_file = BytesIO(response.content)
        return docx_file
    except Exception as e:
        print(f"Ошибка при скачивании {url}: {e}")
        return None


# Основной процесс для обработки списка ссылок
def process_docx_links(links):
    result = {}

    for url in links:
        if not url.lower().endswith('.docx'):
            continue

        print(f"Обрабатывается: {url}")

        # Скачиваем документ
        docx_file = download_docx(url)
        if docx_file:
            # Извлекаем текст
            text_content = extract_text_from_docx(docx_file)
            result[url] = text_content

    return result

# Чтение списка ссылок из файла в MinIO
def read_links_from_minio():
    minio_client, bucket_name = get_minio_client()
    try:
        # Получаем CSV файл из MinIO
        response = minio_client.get_object(bucket_name, "links.csv")
        csv_content = response.read().decode('utf-8')
        response.close()
        response.release_conn()

        # Парсим
        links = [link.strip() for link in csv_content.split('\n') if link.strip()]

        # Фильтруем только DOCX ссылки
        docx_links = [link for link in links if link.lower().endswith('.docx')]

        print(f"Загружено DOCX ссылок из MinIO: {len(docx_links)}")
        return docx_links

    except Exception as e:
        print(f"Ошибка чтения links.csv из MinIO: {e}")
        return []


# Сохранение данных в JSON файл в MinIO
def save_to_minio(data, filename='parsed_docx.json'):
    try:
        minio_client, bucket_name = get_minio_client()

        # Конвертируем данные в JSON
        json_data = json.dumps(data, ensure_ascii=False, indent=2)

        # Сохраняем в MinIO
        minio_client.put_object(
            bucket_name,
            filename,
            io.BytesIO(json_data.encode('utf-8')),
            length=len(json_data.encode('utf-8')),
            content_type="application/json"
        )

        print(f"Результаты сохранены в MinIO: {filename}")

    except Exception as e:
        print(f"Ошибка сохранения в MinIO: {e}")


if __name__ == "__main__":
    # Читаем ссылки из MinIO
    links = read_links_from_minio()

    if not links:
        print("Не найдено ссылок для обработки")
        exit()

    # Обрабатываем документы
    parsed_data = process_docx_links(links)

    # Сохраняем результат в MinIO
    save_to_minio(parsed_data, 'parsed_docx.json')

    # Статистика
    total_files = len(parsed_data)
    total_text_length = sum(len(text) for text in parsed_data.values())
    print(f"Обработано ссылок: {total_files}")
    print(f"Общий объем текста: {total_text_length} символов")